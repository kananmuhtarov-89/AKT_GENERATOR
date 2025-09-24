# -*- coding: utf-8 -*-
import streamlit as st
from io import BytesIO
from datetime import datetime
from typing import List
import re
import pandas as pd

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

st.set_page_config(page_title="AKT (Excel → Word)", page_icon="🧾", layout="centered")

EXCEL_SHEET_DEFAULT = ""  # boş = 1-ci vərəq

PLACEHOLDER_OPTIONS = [
    "NETICELER VE SIYAHI BURA YAZILACAQ.",
    "NETICELER VE SIYAHI BURA YAZILACAQ",
    "NETICƏLƏR VƏ SİYAHI BURA YAZILACAQ.",
    "NETICƏLƏR VƏ SİYAHI BURA YAZILACAQ",
]

BOLD_LABEL_FOR_ALL = True
BOLD_NV = set()  # məsələn {1, 2}

# ========= KÖMƏKÇİ =========
def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())

def _has_placeholder_text(text: str) -> bool:
    t = _norm(text)
    return any(_norm(opt) in t for opt in PLACEHOLDER_OPTIONS)

def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    # Başlıqları lower + diakritikasız müqayisə edək
    def strip_diacritics(x: str) -> str:
        return (
            x.replace("ş", "s").replace("Ş", "S")
             .replace("ı", "i").replace("İ", "I")
             .replace("ə", "e").replace("Ə", "E")
        )
    col_satis = col_siyahi = None
    for c in df.columns:
        lc = strip_diacritics(str(c).lower().strip())
        if "satis" in lc and ("siralama" in lc or "siralamasi" in lc or "siralama" in lc):
            col_satis = c
        if "siyahi" in lc or "siyah" in lc:
            col_siyahi = c
    if col_satis is None or col_siyahi is None:
        raise KeyError("Lazımi sütunlar tapılmadı. Gözlənilən: 'Satış sıralaması' və 'siyahı' (adları yaxın olmalıdır).")
    return df.rename(columns={col_satis: "Satis", col_siyahi: "Nomre"})

def extract_numeric(series: pd.Series) -> pd.Series:
    cleaned = series.astype(str).str.replace(r"\D", "", regex=True)
    return pd.to_numeric(cleaned, errors="coerce")

def build_line_for_one_sale(df: pd.DataFrame, s: int) -> str:
    subset = df.loc[df["Satis"] == s, "Num"].dropna().astype(int)
    nums = sorted(subset.unique().tolist())
    nums_text = ", ".join(str(n) for n in nums) if nums else ""
    return f"{s}-ci NV: {nums_text}"

def set_paragraph_style(p):
    p.paragraph_format.line_spacing = 1.15

def ensure_rpr(run):
    # python-docx guard: ensure run.rPr exists
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    return run._element.rPr

def set_run_arial12(run, bold=False):
    run.font.name = "Arial"
    ensure_rpr(run).rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(12)
    run.bold = bool(bold)

def add_nv_line_to_paragraph(p, line: str, make_label_bold=True, bold_whole=False):
    set_paragraph_style(p)
    m = re.match(r"^(\d+-ci NV:)(\s*)(.*)$", line)
    if not m:
        r = p.add_run(line)
        set_run_arial12(r, bold=bold_whole)
        return
    label, spaces, rest = m.groups()

    r1 = p.add_run(label)
    set_run_arial12(r1, bold=(bold_whole or make_label_bold))

    if spaces:
        r_sp = p.add_run(spaces)
        set_run_arial12(r_sp, bold=bold_whole and not make_label_bold)

    r2 = p.add_run(rest)
    set_run_arial12(r2, bold=bold_whole)

def collect_placeholders(doc: Document):
    found = []
    for p in doc.paragraphs:
        if _has_placeholder_text("".join(r.text for r in p.runs)):
            found.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _has_placeholder_text("".join(r.text for r in p.runs)):
                        found.append(p)
    return found

def fill_placeholders(doc: Document, lines: List[str]):
    targets = collect_placeholders(doc)
    if not targets:
        raise FileNotFoundError("Skeletdə placeholder tapılmadı. Şablonda mətn olaraq `NETICELER VE SIYAHI BURA YAZILACAQ` yazısı olmalıdır.")
    if len(targets) == 1:
        p = targets[0]
        p.text = ""
        first = True
        for line in lines:
            if not first:
                p.add_run().add_break()
            first = False
            nv_no = None
            if "-ci NV:" in line:
                try:
                    nv_no = int(line.split("-ci NV:")[0])
                except:
                    nv_no = None
            bold_whole = (nv_no in BOLD_NV) if nv_no is not None else False
            add_nv_line_to_paragraph(p, line, make_label_bold=BOLD_LABEL_FOR_ALL, bold_whole=bold_whole)
    else:
        count = min(len(targets), len(lines))
        for i in range(count):
            p = targets[i]
            p.text = ""
            line = lines[i]
            nv_no = None
            if "-ci NV:" in line:
                try:
                    nv_no = int(line.split("-ci NV:")[0])
                except:
                    nv_no = None
            bold_whole = (nv_no in BOLD_NV) if nv_no is not None else False
            add_nv_line_to_paragraph(p, line, make_label_bold=BOLD_LABEL_FOR_ALL, bold_whole=bold_whole)

def build_output_name(sales_list: List[int]) -> str:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    tag = "NV-" + "-".join(str(s) for s in sales_list) if sales_list else "NV"
    return f"AKT_{stamp}__{tag}.docx"

# ============== UI ==============
st.title("🧾 AKT Doldurma (Excel → Word)")
st.caption("Excel-dən NV siyahılarını oxuyub Word şablonuna yerləşdirir — **tək .docx çıxış**.")

colf1, colf2 = st.columns([1,1])
with colf1:
    excel_file = st.file_uploader("Excel (.xlsx) – məlumat", type=["xlsx"])
with colf2:
    docx_file = st.file_uploader("Word (.docx) – şablon (placeholder mətn olsun)", type=["docx"])

sheet = st.text_input("Vərəq adı (boş saxla: 1-ci vərəq)", value=EXCEL_SHEET_DEFAULT)
sales_raw = st.text_input("NV satış nömrələri (vergüllə)", value="", placeholder="məs: 1,2,3")

go = st.button("AKT yarad və endir")

if go:
    if not excel_file or not docx_file:
        st.error("Həm Excel (.xlsx), həm də Word (.docx) faylı yükləməlisiniz.")
        st.stop()
    try:
        # parse sales
        try:
            sales_list = [int(x.strip()) for x in sales_raw.split(",") if x.strip()]
        except Exception:
            st.error("NV satış nömrələri sırf rəqəm olmalıdır (məs: 1,2,3).")
            st.stop()
        if not sales_list:
            st.error("NV siyahısı boşdur.")
            st.stop()

        # read excel
        if sheet.strip():
            df = pd.read_excel(excel_file, sheet_name=sheet.strip(), dtype=object, engine="openpyxl")
        else:
            df = pd.read_excel(excel_file, dtype=object, engine="openpyxl")
        df = normalize_columns(df)[["Satis", "Nomre"]].copy()
        df["Satis"] = pd.to_numeric(df["Satis"], errors="coerce").ffill().astype("Int64")
        df["Num"] = extract_numeric(df["Nomre"])

        # build lines
        lines = [build_line_for_one_sale(df, s) for s in sales_list]

        # fill docx
        doc = Document(docx_file)
        fill_placeholders(doc, lines)

        out_name = build_output_name(sales_list)
        buf = BytesIO()
        doc.save(buf); buf.seek(0)

        st.success(f"Hazırdır: {out_name}")
        st.download_button("Docx endir", data=buf.getvalue(), file_name=out_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(f"Xəta: {e}")
